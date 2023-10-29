import pandas as pd
from sklearn.preprocessing import StandardScaler
from sklearn.cluster import KMeans
from sklearn.metrics import silhouette_score
import matplotlib.pyplot as plt
from docx import Document
import logging

class CustomerSegmentation:
    """
    Perform customer segmentation based on various features.

    Attributes:
    - file_path: The path to the data file.
    - n_clusters: The number of clusters to form.
    - clustering_algorithm: The algorithm to use for clustering. Defaults to 'kmeans'.
    """

    def __init__(self, file_path, clustering_algorithm='kmeans', n_clusters=4):
        """
        Initialize the CustomerSegmentation class.

        Parameters:
        - file_path: str, path to the data file.
        - clustering_algorithm: str, clustering algorithm to use. Defaults to 'kmeans'.
        - n_clusters: int, number of clusters to form. Defaults to 4.
        """
        # Setting up logging
        logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
        self.file_path = file_path
        self.n_clusters = n_clusters
        self.clustering_algorithm = clustering_algorithm

    def validate_data(self, data):
        """
        Validate the input data against required columns.

        Parameters:
        - data: DataFrame, the input data.

        Raises:
        - ValueError: If any required column is missing.

        Logs:
        - Info: When data validation passes.
        """
        required_columns = [
            'Age', 'Gender', 'Location', 'Subscription Status',
            'Payment Method', 'Membership Level', 'Return History',
            'Browsing History', 'Inseam Size', 'Waist Size', 'Shoe Size'
        ]
        for col in required_columns:
            if col not in data.columns:
                raise ValueError(f"Column {col} is missing from the dataset.")
        logging.info('Data validation passed.')


    @staticmethod
    def save_summary_to_word(summary, file_name='cluster_summary.docx'):
        """
        Save the cluster summary to a Word document.

        Parameters:
        - summary: DataFrame, the summary data to save.
        - file_name: str, the name of the Word document to create. Defaults to 'cluster_summary.docx'.

        Logs:
        - Info: When the summary is successfully saved.
        - Error: When an error occurs while saving.
        """
        try:
            doc = Document()
            doc.add_heading('Cluster Summary', level=1)

            # Creating a table in Word
            table = doc.add_table(rows=1, cols=summary.shape[1])
            table.autofit = True

            # Adding headers
            hdr_cells = table.rows[0].cells
            for idx, col in enumerate(summary.columns):
                hdr_cells[idx].text = str(col)

            # Adding data rows
            for index, row in summary.iterrows():
                cells = table.add_row().cells
                for idx, value in enumerate(row):
                    cells[idx].text = str(value)

            doc.save(file_name)
            logging.info(f'Summary saved to {file_name}')

        except FileNotFoundError:
            logging.error('File not found.')
        except PermissionError:
            logging.error('Permission denied.')
        except Exception as e:
            logging.error(f'An unknown error occurred: {e}')

    def export_summary_to_csv(self, summary, file_name='cluster_summary.csv'):
        """
        Export the cluster summary to a CSV file.

        Parameters:
        - summary: DataFrame, the summary data to export.
        - file_name: str, the name of the CSV file to create. Defaults to 'cluster_summary.csv'.

        Logs:
        - Info: When the summary is successfully exported.
        - Error: When an error occurs while exporting.
        """
        try:
            # Check if 'summary' is a DataFrame
            if not isinstance(summary, pd.DataFrame):
                raise TypeError("Input summary must be a DataFrame.")
            
            summary.to_csv(file_name)
            logging.info(f'Summary exported to {file_name}')

        except FileNotFoundError:
            logging.error('File not found.')
        except PermissionError:
            logging.error('Permission denied.')
        except TypeError as te:
            logging.error(f'Type error: {te}')
        except Exception as e:
            logging.error(f'An unknown error occurred while exporting summary to CSV: {e}')

    def load_and_prepare_data(self):
        """
        Load and prepare customer data for segmentation.
        
        Returns:
        DataFrame: A DataFrame containing only the features needed for segmentation.
        
        Raises:
        FileNotFoundError: If the file doesn't exist.
        Exception: For any other error.
        """
        try:
            customer_data = pd.read_excel(self.file_path)
            self.validate_data(customer_data)
            segmentation_features = [
                'Age', 'Gender', 'Location', 'Subscription Status',
                'Payment Method', 'Membership Level', 'Return History',
                'Browsing History', 'Inseam Size', 'Waist Size', 'Shoe Size'
            ]
            return customer_data[segmentation_features]
        except FileNotFoundError:
            logging.error('Data file not found.')
            raise
        except Exception as e:
            logging.error(f'Error in data loading and preparation: {e}')
            raise

    def encode_and_scale_data(self, data):
        """
        Encode categorical variables and scale numerical features.
        
        Parameters:
        data (DataFrame): Data to be encoded and scaled.
        
        Returns:
        ndarray: Scaled and encoded features.
        
        Raises:
        Exception: For any error during encoding or scaling.
        """
        try:
            encoded_data = pd.get_dummies(data, drop_first=True)
            scaler = StandardScaler()
            return scaler.fit_transform(encoded_data.select_dtypes(include=['number']))
        except Exception as e:
            logging.error(f'Error in encoding and scaling data: {e}')
            raise

    def apply_clustering(self, scaled_data):
        """
        Apply K-Means clustering to the scaled data.

        Parameters:
        - scaled_data (ndarray): The scaled feature set.

        Returns:
        - ndarray: Cluster labels for each data point.

        Raises:
        - Exception: For any error during clustering.
        """
        try:
            kmeans = KMeans(n_clusters=self.n_clusters, random_state=42, n_init=10)
            kmeans.fit(scaled_data)
            return kmeans.labels_
        except Exception as e:
            logging.error(f'Error during clustering: {e}')
            raise

    def calculate_cluster_summary(self, data, labels):
        """
        Calculate the mean of each feature for each cluster.

        Parameters:
        - data (DataFrame): The original data.
        - labels (ndarray): Cluster labels.

        Returns:
        - DataFrame: Summary statistics for each cluster.

        Raises:
        - Exception: For any error during summary calculation.
        """
        try:
            data['Cluster'] = labels
            numeric_columns = data.select_dtypes(include=['number'])
            return numeric_columns.groupby('Cluster').mean().round(2)
        except Exception as e:
            logging.error(f'Error while calculating cluster summary: {e}')
            raise

    def display_cluster_summary(self, cluster_summary):
        """
        Display the cluster summary.

        Parameters:
        - cluster_summary (DataFrame): The summary to display.
        """
        print(cluster_summary)

    def calculate_and_display_silhouette_score(self, scaled_data, labels):
        """
        Calculate and display the silhouette score for clustering.

        Parameters:
        - scaled_data (ndarray): The scaled feature set.
        - labels (ndarray): Cluster labels.

        Raises:
        - Exception: For any error during silhouette score calculation.
        """
        try:
            silhouette_avg = silhouette_score(scaled_data, labels)
            print(f'Silhouette Score: {silhouette_avg:.2f}')
        except Exception as e:
            logging.error(f'Error during silhouette score calculation: {e}')
            raise

    def visualize_and_save_clusters(self, data, labels):
        """
        Visualize the clusters in a scatter plot and save it as a PNG file.

        Parameters:
        - data (DataFrame): The original data.
        - labels (ndarray): Cluster labels.

        Raises:
        - Exception: For any error during plotting or saving the plot.
        """
        try:
            plt.scatter(data['Age'], data['Inseam Size'], c=labels)
            plt.xlabel('Age')
            plt.ylabel('Inseam Size')
            plt.title('Clusters based on Age and Inseam Size')
            plt.savefig('clusters_plot.png')
            plt.show()
        except Exception as e:
            logging.error(f'Error during visualization: {e}')
            raise

    def perform_segmentation(self):
        """
        Perform the entire customer segmentation pipeline.

        Steps include:
        1. Data Preparation
        2. Feature Scaling and Encoding
        3. Clustering
        4. Summary Calculation
        5. Display and Save Results

        Logs:
        - Info: When the segmentation is successful and results are saved.
        - Error: When any error occurs during the process.
        """
        try:
            prepared_data = self.load_and_prepare_data()
        except Exception as e:
            logging.error(f'Error during data preparation: {e}')
            return

        try:
            scaled_data = self.encode_and_scale_data(prepared_data)
        except Exception as e:
            logging.error(f'Error during data scaling and encoding: {e}')
            return

        try:
            cluster_labels = self.apply_clustering(scaled_data)
        except Exception as e:
            logging.error(f'Error during clustering: {e}')
            return

        try:
            cluster_summary = self.calculate_cluster_summary(prepared_data, cluster_labels)
        except Exception as e:
            logging.error(f'Error during summary calculation: {e}')
            return

        try:
            self.display_cluster_summary(cluster_summary)
            self.calculate_and_display_silhouette_score(scaled_data, cluster_labels)
            self.save_summary_to_word(cluster_summary)
            self.export_summary_to_csv(cluster_summary)
            self.visualize_and_save_clusters(prepared_data, cluster_labels)

            logging.info('Customer segmentation successful and results saved.')
        except Exception as e:
            logging.error(f'An error occurred during the customer segmentation process: {e}')


# Usage example
file_path = 'file_path' # Change this to your file path
segmentation = CustomerSegmentation(file_path)
segmentation.perform_segmentation()
